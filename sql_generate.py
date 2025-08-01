import re
import os
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

# 设置重复表头
def set_repeat_table_header(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)
    
# 设置表头不换行
def set_no_wrap(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    noWrap = OxmlElement('w:noWrap')
    tcPr.append(noWrap)

# 设置字体
def set_font(cell, name='宋体', size=12):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
            run.font.size = Pt(size)
            # 设置语言为中文，防止拼写标红
            rPr = run._element.get_or_add_rPr()
            lang = OxmlElement('w:lang')
            lang.set(qn('w:val'), 'zh-CN')
            rPr.append(lang)

# 设置背景色
def set_cell_background(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

# 禁止拆除行
def prevent_row_split(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    cant_split = OxmlElement('w:cantSplit')
    trPr.append(cant_split)

def center_cell(cell):
    # 设置垂直居中
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 设置水平居中
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 大小驼峰转换
def to_camel_case(name: str) -> str:
    parts = name.lower().split('_')
    return parts[0] + ''.join(p.capitalize() for p in parts[1:])

# 表结构和注释 SQL
with open('sqltxt.sql', encoding='utf-8') as f:
    create_sql = f.read()

def set_cell_padding(cell, top=100, bottom=100, left=100, right=100):
    """
    设置单元格内边距，单位是twips（1pt=20twips）
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcMar = tcPr.find(qn('w:tcMar'))
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)

    def set_margin(name, value):
        node = tcMar.find(qn(f'w:{name}'))
        if node is None:
            node = OxmlElement(f'w:{name}')
            tcMar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')  # dxa表示twips单位
    set_margin('top', top)
    set_margin('bottom', bottom)
    set_margin('left', left)
    set_margin('right', right)
print('ε=( o｀ω′)ノ进行文本提取字段...')
# 1) 提取字段
field_pattern = re.compile(
    r'^\s*"(?P<name>\w+)"\s+(?P<type>\w+(?:\(\d+(?:,\s*\d+)?\))?)\s*(?P<extra>[^,\n]*)(?:,|\n)',
    re.MULTILINE | re.IGNORECASE
)
fields = field_pattern.findall(create_sql)
print('(๑•̀ㅂ•́)و✧字段解析完成！')
print('')
# 2) 提取主键字段列表
# 匹配 PRIMARY KEY (...) 里面的字段名，允许多字段复合主键
print('ε=( o｀ω′)ノ进行复合主键提取...')
pk_pattern = re.compile(r'PRIMARY KEY\s*\(([^)]+)\)', re.IGNORECASE)
pk_match = pk_pattern.search(create_sql)
if pk_match:
    pk_fields_str = pk_match.group(1)
    # 去除双引号，拆分字段名
    pk_fields = [f.strip().strip('"') for f in pk_fields_str.split(',')]
else:
    pk_fields = []
print('(๑•̀ㅂ•́)و✧主键提取完成！')
print('')
print('ε=( o｀ω′)ノ进行注释提取...')
# 提取注释
comment_pattern = re.compile(r'COMMENT ON COLUMN "public"\."[^"]+"\."(\w+)" IS \'([^\']*)\'', re.IGNORECASE)
comments = dict(comment_pattern.findall(create_sql))
print('(๑•̀ㅂ•́)و✧注释提取完成！')
print('')
print('ε=( o｀ω′)ノ开始生成Word结构内容...')
# 创建文档
doc = Document()
# 创建表格
table = doc.add_table(rows=1, cols=6)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

# 表头宽度设置，这里定义每列宽度，单位是厘米
col_widths = [16, 16, 16, 16, 16, 16]
# 设置表头文本
headers = ['字段中文名称', '字段名称', '字段类型', '长度', '备注', '对应代码名称']
hdr_cells = table.rows[0].cells
for i, text in enumerate(headers):
    hdr_cells[i].text = text
    set_font(hdr_cells[i], name='宋体', size=12)
    set_cell_background(hdr_cells[i], 'cecece')  # 设置灰色背景
    set_no_wrap(hdr_cells[i])
    center_cell(hdr_cells[i])  # 垂直居中
    #设置表头宽高
    set_cell_padding(hdr_cells[i], top=100, bottom=100, left=0, right=0)
    hdr_cells[i].width = Cm(col_widths[i])
# 设置表头重复
set_repeat_table_header(table.rows[0])

# 添加字段行
for field_name, field_type, rest in fields:
    full_def = (field_type + rest).upper()
    row_cells = table.add_row().cells
    prevent_row_split(table.rows[-1])  # <== 禁止当前行被分页断开
    comment = comments.get(field_name, field_name)  # 如果没有注释，用字段名代替
    row_cells[i].width = Cm(col_widths[i])
    row_cells[i].height = Cm(1.1)
    # 处理长度
    length_match = re.search(r'\((\d+)(?:,\d+)?\)', field_type)
    if length_match:
        length = length_match.group(1)
    else:
        ft = field_type.lower()
        if 'int8' in ft:
            length = '19'
        elif 'int4' in ft:
            length = '10'
        elif 'int2' in ft:
            length = '5'
        elif 'numeric' in ft:
            match = re.search(r'numeric\((\d+),\s*\d+\)', ft)
            length = match.group(1) if match else '0'
        elif 'float8' in ft or 'double precision' in ft:
            length = '17'  # 双精度浮点大约可表示17位十进制数
        elif 'float4' in ft or 'real' in ft:
            length = '6'   # 单精度浮点大约6位有效数字
        elif 'varchar' in ft or 'character varying' in ft:
            match = re.search(r'\((\d+)\)', ft)
            length = match.group(1) if match else '不限'
        elif 'char' in ft or 'character' in ft:
            match = re.search(r'\((\d+)\)', ft)
            length = match.group(1) if match else '不限'
        elif 'text' in ft or 'json' in ft or 'bytea' in ft:
            length = '不限'
        elif 'boolean' in ft or 'bool' in ft:
            length = '1'
        elif 'date' in ft:
            length = '10'  # YYYY-MM-DD
        elif 'time' in ft:
            length = '8'  # HH:MM:SS
        elif 'timestamp' in ft:
            length = '26'  # 包括小数点秒
        elif 'uuid' in ft:
            length = '36'
        elif 'inet' in ft or 'cidr' in ft:
            length = '43'
        else:
            length = '0'
    # 处理注释
    remark = ''
    if field_name in pk_fields:
        remark = '主键，'
    if 'NOT NULL' in full_def.upper():
        remark += '不允许为空'
    else:
        remark += '允许为空'
    # 处理代码名称
    codeName = to_camel_case(field_name)
    values = [
        comment,
        field_name,
        field_type.upper(),
        length,
        remark , 
        codeName
    ]
    for i, val in enumerate(values):
        row_cells[i].text = val
        set_font(row_cells[i], name='宋体', size=12)
        center_cell(row_cells[i])  # 垂直居中
# 保存路径
output_path = os.path.abspath("./generate-doc.docx")
doc.save(output_path)

print("(๑•̀ㅂ•́)و✧文档已生成成功，保存路径：", output_path)
